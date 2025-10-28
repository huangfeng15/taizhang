"""
增强版Word报告导出模块
提供更详实的文字内容和专业的报告格式
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def export_to_word_enhanced(report_data, file_path):
    """
    导出增强版Word报告，包含大量文字内容
    
    Args:
        report_data: 报表数据字典
        file_path: 导出文件路径
    
    Returns:
        str: 导出的文件路径
    """
    
    doc = Document()
    
    # 设置文档默认字体为中文
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.styles['Normal'].font.size = Pt(12)
    
    # 判断报告类型
    report_type = report_data.get('report_type', 'monthly')
    is_project_report = report_type == 'project'
    
    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(report_data.get('title', '工作报表'))
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 102)
    title_run.font.name = '黑体'
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    
    doc.add_paragraph()
    
    # 基本信息
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info_para.add_run(
        f"统计周期：{report_data['period_start']} 至 {report_data['period_end']}\n"
        f"生成时间：{report_data['generated_at'].strftime('%Y年%m月%d日 %H:%M:%S')}"
    )
    info_run.font.size = Pt(11)
    info_run.font.color.rgb = RGBColor(102, 102, 102)
    
    doc.add_paragraph()
    
    # ==================== 一、工作概况 ====================
    summary_heading = doc.add_heading('一、工作概况', level=1)
    summary_heading.runs[0].font.size = Pt(16)
    summary_heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    summary_heading.runs[0].font.name = '黑体'
    summary_heading.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    
    summary = report_data['summary']
    
    if is_project_report:
        # 项目报告的工作概况
        intro_para = doc.add_paragraph()
        intro_para.add_run(
            "本项目自启动以来，在各级领导的高度重视和大力支持下，项目采购与成本管理工作始终坚持规范化、精细化、信息化的管理理念，"
            "严格遵守国家法律法规和相关管理制度要求，扎实推进各项业务工作。项目团队认真履职尽责，主动作为，"
            "在采购管理、合同管理、资金支付、结算管理等各个环节都建立了完善的管理体系和工作机制，"
            "确保了项目建设的有序推进和管理目标的顺利实现。截至报告期末，项目在各个业务领域均取得了显著进展，"
            "为项目的全面完成奠定了坚实基础。"
        ).font.size = Pt(12)
        
        doc.add_paragraph()
        
        # 数据段落
        data_intro = doc.add_paragraph()
        data_intro.add_run(
            "从项目全生命周期的业务数据来看，各项指标充分反映了项目管理的规范性和有效性。"
            "在采购环节，项目严格执行采购管理制度，规范开展采购活动。项目已完成采购项目"
        ).font.size = Pt(12)
        
        proc_count = data_intro.add_run(f" {summary['total_procurement_count']} ")
        proc_count.font.size = Pt(12)
        proc_count.font.bold = True
        proc_count.font.color.rgb = RGBColor(192, 0, 0)
        
        data_intro.add_run(
            "个，采购工作覆盖了项目建设的各个领域和环节，充分满足了项目建设需求。项目中标总金额"
        ).font.size = Pt(12)
        
        winning_amt = data_intro.add_run(f" {summary['total_winning_amount']:,.2f} ")
        winning_amt.font.size = Pt(12)
        winning_amt.font.bold = True
        winning_amt.font.color.rgb = RGBColor(192, 0, 0)
        
        data_intro.add_run(
            "元，通过科学的成本控制和充分的市场竞争，实现了采购成本的有效控制。"
        ).font.size = Pt(12)
        
        doc.add_paragraph()
        contract_para = doc.add_paragraph()
        contract_para.add_run(
            "在合同管理环节，项目建立了完善的合同管理体系，从合同签订、履约管理到验收评价，"
            "每个环节都有明确的工作标准和操作规程。项目累计签订合同"
        ).font.size = Pt(12)
        
        cont_count = contract_para.add_run(f" {summary['total_contract_count']} ")
        cont_count.font.size = Pt(12)
        cont_count.font.bold = True
        cont_count.font.color.rgb = RGBColor(192, 0, 0)
        
        contract_para.add_run(
            "份，合同总金额"
        ).font.size = Pt(12)
        
        cont_amt = contract_para.add_run(f" {summary['total_contract_amount']:,.2f} ")
        cont_amt.font.size = Pt(12)
        cont_amt.font.bold = True
        cont_amt.font.color.rgb = RGBColor(192, 0, 0)
        
        contract_para.add_run(
            "元。所有合同均经过严格的审核把关，合同条款完备，权利义务明确，"
            "有效保障了项目利益和合同双方的合法权益。"
        ).font.size = Pt(12)
        
        doc.add_paragraph()
        payment_para = doc.add_paragraph()
        payment_para.add_run(
            "在资金支付环节，项目建立了严格的资金支付管理制度，确保资金支付的及时性、准确性和合规性。"
            "项目累计处理付款业务"
        ).font.size = Pt(12)
        
        pay_count = payment_para.add_run(f" {summary['total_payment_count']} ")
        pay_count.font.size = Pt(12)
        pay_count.font.bold = True
        pay_count.font.color.rgb = RGBColor(192, 0, 0)
        
        payment_para.add_run(
            "笔，累计付款金额"
        ).font.size = Pt(12)
        
        pay_amt = payment_para.add_run(f" {summary['total_payment_amount']:,.2f} ")
        pay_amt.font.size = Pt(12)
        pay_amt.font.bold = True
        pay_amt.font.color.rgb = RGBColor(192, 0, 0)
        
        payment_para.add_run(
            "元。资金支付严格按照合同约定和资金计划执行，既保障了项目建设的资金需求，"
            "又实现了资金的合理调配和高效使用。"
        ).font.size = Pt(12)
        
        doc.add_paragraph()
        settlement_para = doc.add_paragraph()
        settlement_para.add_run(
            "在结算管理环节，项目高度重视结算工作的规范性和及时性，建立了完善的结算管理流程。"
            "项目累计完成结算"
        ).font.size = Pt(12)
        
        settle_count = settlement_para.add_run(f" {summary['total_settlement_count']} ")
        settle_count.font.size = Pt(12)
        settle_count.font.bold = True
        settle_count.font.color.rgb = RGBColor(192, 0, 0)
        
        settlement_para.add_run(
            "笔，结算总金额"
        ).font.size = Pt(12)
        
        settle_amt = settlement_para.add_run(f" {summary['total_settlement_amount']:,.2f} ")
        settle_amt.font.size = Pt(12)
        settle_amt.font.bold = True
        settle_amt.font.color.rgb = RGBColor(192, 0, 0)
        
        settlement_para.add_run(
            "元。结算工作的及时完成，不仅为项目的财务核算提供了准确依据，"
            "也为项目的最终验收和评价工作奠定了坚实基础。"
        ).font.size = Pt(12)
        
        # 项目执行进度分析
        if summary['total_contract_amount'] > 0:
            doc.add_paragraph()
            progress_para = doc.add_paragraph()
            payment_ratio = (summary['total_payment_amount'] / summary['total_contract_amount']) * 100
            settlement_ratio = (summary['total_settlement_count'] / summary['total_contract_count'] * 100) if summary['total_contract_count'] > 0 else 0
            
            progress_para.add_run(
                "从项目整体执行情况来看，项目付款进度为"
            ).font.size = Pt(12)
            
            pay_ratio_run = progress_para.add_run(f" {payment_ratio:.1f}% ")
            pay_ratio_run.font.size = Pt(12)
            pay_ratio_run.font.bold = True
            pay_ratio_run.font.color.rgb = RGBColor(0, 176, 80)
            
            progress_para.add_run(
                "，结算完成比例为"
            ).font.size = Pt(12)
            
            settle_ratio_run = progress_para.add_run(f" {settlement_ratio:.1f}% ")
            settle_ratio_run.font.size = Pt(12)
            settle_ratio_run.font.bold = True
            settle_ratio_run.font.color.rgb = RGBColor(0, 176, 80)
            
            if payment_ratio >= 80:
                progress_para.add_run(
                    "。项目已进入收尾阶段，资金支付和合同履约情况良好，项目整体运行顺利，"
                    "资金支付与合同执行保持同步推进。下一步应重点关注剩余合同的履约和结算工作，"
                    "确保项目按期全面完成。"
                ).font.size = Pt(12)
            elif payment_ratio >= 60:
                progress_para.add_run(
                    "。项目进展顺利，各项工作有序推进，资金支付与合同执行保持良好的同步性。"
                    "后续应继续加强合同履约管理，及时办理资金支付，确保项目建设不受资金制约。"
                ).font.size = Pt(12)
            else:
                progress_para.add_run(
                    "。项目处于建设期，各项工作正在积极推进中。应继续加强项目管理，"
                    "优化工作流程，提高工作效率，确保项目按计划推进。"
                ).font.size = Pt(12)
                
    else:
        # 时间区间报告的工作概况
        intro_para = doc.add_paragraph()
        intro_para.add_run(
            "本报告期内，在各级领导的正确领导和各部门的密切配合下，项目采购与成本管理工作紧紧围绕年度工作目标，"
            "坚持问题导向和目标导向相结合，扎实推进各项工作任务。我们始终把规范管理、提升效率、控制成本作为工作重点，"
            "通过完善制度体系、优化工作流程、强化过程管控、加强队伍建设等措施，不断提升管理水平和服务质量。"
            "期间，各业务领域工作有序开展，管理质量持续提升，各项指标完成情况良好，"
            "为全年工作目标的实现奠定了坚实基础。"
        ).font.size = Pt(12)
        
        doc.add_paragraph()
        
        # 数据段落
        data_intro = doc.add_paragraph()
        data_intro.add_run(
            "从本期业务开展情况和主要数据指标来看，各项工作均取得了积极进展。"
            "采购业务方面，本期采购工作任务饱满，组织有力。共完成采购项目"
        ).font.size = Pt(12)
        
        proc_count = data_intro.add_run(f" {summary['total_procurement_count']} ")
        proc_count.font.size = Pt(12)
        proc_count.font.bold = True
        proc_count.font.color.rgb = RGBColor(192, 0, 0)
        
        data_intro.add_run(
            "个，采购工作涵盖了项目建设的多个领域，充分保障了项目建设需求。本期中标总金额"
        ).font.size = Pt(12)
        
        winning_amt = data_intro.add_run(f" {summary['total_winning_amount']:,.2f} ")
        winning_amt.font.size = Pt(12)
        winning_amt.font.bold = True
        winning_amt.font.color.rgb = RGBColor(192, 0, 0)
        
        data_intro.add_run(
            "元，采购规模合理，价格水平适中，充分体现了市场竞争原则和成本控制要求。"
        ).font.size = Pt(12)
        
        doc.add_paragraph()
        contract_para = doc.add_paragraph()
        contract_para.add_run(
            "合同管理方面，本期合同签订工作扎实推进，合同管理水平不断提升。本期签订合同"
        ).font.size = Pt(12)
        
        cont_count = contract_para.add_run(f" {summary['total_contract_count']} ")
        cont_count.font.size = Pt(12)
        cont_count.font.bold = True
        cont_count.font.color.rgb = RGBColor(192, 0, 0)
        
        contract_para.add_run(
            "份，合同总金额"
        ).font.size = Pt(12)
        
        cont_amt = contract_para.add_run(f" {summary['total_contract_amount']:,.2f} ")
        cont_amt.font.size = Pt(12)
        cont_amt.font.bold = True
        cont_amt.font.color.rgb = RGBColor(192, 0, 0)
        
        contract_para.add_run(
            