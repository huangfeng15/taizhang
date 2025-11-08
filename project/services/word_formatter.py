"""
统一的Word文档格式化器
遵循OCP原则：通过模板类型扩展功能，无需修改现有代码
遵循DRY原则：统一样式设置，避免重复
合并了word_exporter.py和comprehensive_word_exporter.py的功能
"""
from typing import Dict, Any, List
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


class WordFormatter:
    """
    Word文档格式化器
    支持多种报表模板：standard（标准）、professional（专业）、comprehensive（综合）
    """

    def __init__(self, template_type: str = 'standard'):
        """
        初始化Word格式化器

        Args:
            template_type: 模板类型 ('standard' | 'professional' | 'comprehensive')
        """
        self.template_type = template_type
        self.doc = Document()
        self._setup_document()

    def _setup_document(self):
        """设置文档样式（DRY - 统一样式设置）"""
        # 设置默认字体
        style = self.doc.styles['Normal']
        style.font.name = '宋体'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        style.font.size = Pt(12)

        # 专业和综合报告使用1.5倍行距
        if self.template_type in ['professional', 'comprehensive']:
            style.paragraph_format.line_spacing = 1.5

        # 设置页面边距
        for section in self.doc.sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(3.17)
            section.right_margin = Cm(3.17)

    def format_report(self, report_data: Dict[str, Any], file_path: str) -> str:
        """
        格式化报表数据为Word文档

        Args:
            report_data: 报表数据字典
            file_path: 输出文件路径

        Returns:
            str: 生成的文件路径
        """
        # 根据template_type选择不同的格式化策略（OCP原则）
        if self.template_type == 'standard':
            self._format_standard(report_data)
        elif self.template_type == 'professional':
            self._format_professional(report_data)
        elif self.template_type == 'comprehensive':
            self._format_comprehensive(report_data)

        self.doc.save(file_path)
        return file_path

    def _format_standard(self, data: Dict[str, Any]):
        """标准格式（原report_generator.py的逻辑）"""
        meta = data.get('meta', {})
        summary = data.get('summary', {})

        # 封面
        self._add_simple_cover(meta)
        self.doc.add_page_break()

        # 执行摘要
        self._add_section_title('执行摘要', level=1)
        self._add_summary_content(summary)

        # 业务数据
        if data.get('procurement'):
            self._add_section_title('采购业务', level=1)
            self._add_procurement_section(data['procurement'])

        if data.get('contract'):
            self._add_section_title('合同管理', level=1)
            self._add_contract_section(data['contract'])

        if data.get('payment'):
            self._add_section_title('付款业务', level=1)
            self._add_payment_section(data['payment'])

        if data.get('settlement'):
            self._add_section_title('结算业务', level=1)
            self._add_settlement_section(data['settlement'])

    def _format_professional(self, data: Dict[str, Any]):
        """专业格式（原word_exporter.py的逻辑）"""
        meta = data.get('meta', {})

        # 封面
        self._add_professional_cover(meta)
        self.doc.add_page_break()

        # 目录
        self._add_table_of_contents()
        self.doc.add_page_break()

        # 一、执行摘要
        self._add_section_title('一、执行摘要', level=1)
        self._add_summary_content(data.get('summary', {}))
        self.doc.add_paragraph()

        # 二、项目概览
        if data.get('projects_overview'):
            self._add_section_title('二、项目概览', level=1)
            self._add_projects_overview(data['projects_overview'])
            self.doc.add_paragraph()

        # 三、采购业务分析
        if data.get('procurement'):
            self._add_section_title('三、采购业务分析', level=1)
            self._add_procurement_section(data['procurement'])
            self.doc.add_paragraph()

        # 四、合同管理分析
        if data.get('contract'):
            self._add_section_title('四、合同管理分析', level=1)
            self._add_contract_section(data['contract'])
            self.doc.add_paragraph()

        # 五、付款业务分析
        if data.get('payment'):
            self._add_section_title('五、付款业务分析', level=1)
            self._add_payment_section(data['payment'])
            self.doc.add_paragraph()

        # 六、结算业务分析
        if data.get('settlement'):
            self._add_section_title('六、结算业务分析', level=1)
            self._add_settlement_section(data['settlement'])
            self.doc.add_paragraph()

        # 七、归档监控
        if data.get('archive_monitoring'):
            self._add_section_title('七、归档监控', level=1)
            self._add_archive_monitoring(data['archive_monitoring'])
            self.doc.add_paragraph()

        # 八、数据完整性分析
        if data.get('completeness'):
            self._add_section_title('八、数据完整性分析', level=1)
            self._add_completeness_analysis(data['completeness'])
            self.doc.add_paragraph()

        # 九、业务排名
        if data.get('ranking'):
            self._add_section_title('九、业务排名', level=1)
            self._add_ranking_analysis(data['ranking'])
            self.doc.add_paragraph()

        # 十、管理建议
        if data.get('recommendations'):
            self._add_section_title('十、管理建议', level=1)
            self._add_recommendations(data['recommendations'])

    def _format_comprehensive(self, data: Dict[str, Any]):
        """综合格式（原comprehensive_word_exporter.py的逻辑）"""
        meta = data.get('meta', {})

        # 封面
        self._add_comprehensive_cover(meta)
        self.doc.add_page_break()

        # 目录
        self._add_comprehensive_toc()
        self.doc.add_page_break()

        # 第一部分：执行摘要
        self._add_section_title('第一部分  执行摘要', level=1)
        self._add_summary_content(data.get('executive_summary', {}))
        self.doc.add_page_break()

        # 第二部分：组织概览
        if data.get('organizational_overview'):
            self._add_section_title('第二部分  组织概览', level=1)
            self._add_organizational_overview(data['organizational_overview'])
            self.doc.add_page_break()

        # 第三部分：项目分析
        if data.get('projects_analysis'):
            self._add_section_title('第三部分  项目深度分析', level=1)
            self._add_projects_deep_analysis(data['projects_analysis'])
            self.doc.add_page_break()

        # 第四-七部分：业务综合分析
        self._add_business_comprehensive_sections(data)

        # 第八部分：财务分析
        if data.get('financial_analysis'):
            self._add_section_title('第八部分  财务分析', level=1)
            self._add_financial_analysis(data['financial_analysis'])
            self.doc.add_page_break()

        # 第九部分：建议与展望
        if data.get('recommendations'):
            self._add_section_title('第九部分  建议与展望', level=1)
            self._add_recommendations(data['recommendations'])

    # ========== 封面相关方法 ==========

    def _add_simple_cover(self, meta: Dict):
        """添加简单封面"""
        for _ in range(5):
            self.doc.add_paragraph()

        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run(meta.get('report_title', '工作报告'))
        title_run.font.size = Pt(22)
        title_run.font.bold = True
        title_run.font.name = '黑体'
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

        self.doc.add_paragraph()

        # 时间范围
        period = self.doc.add_paragraph()
        period.alignment = WD_ALIGN_PARAGRAPH.CENTER
        period_text = f"{meta.get('period_start', '')} 至 {meta.get('period_end', '')}"
        period_run = period.add_run(period_text)
        period_run.font.size = Pt(14)

    def _add_professional_cover(self, meta: Dict):
        """添加专业封面"""
        for _ in range(5):
            self.doc.add_paragraph()

        # 报告单位
        unit = self.doc.add_paragraph()
        unit.alignment = WD_ALIGN_PARAGRAPH.CENTER
        unit_run = unit.add_run(meta.get('reporting_unit', '项目采购与成本管理部门'))
        unit_run.font.size = Pt(22)
        unit_run.font.bold = True
        unit_run.font.color.rgb = RGBColor(0, 32, 96)
        unit_run.font.name = '黑体'
        unit_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

        self.doc.add_paragraph()

        # 报告标题
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run(meta.get('report_title', '工作报告'))
        title_run.font.size = Pt(26)
        title_run.font.bold = True
        title_run.font.name = '黑体'
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

        self.doc.add_paragraph()

        # 报告范围
        scope = self.doc.add_paragraph()
        scope.alignment = WD_ALIGN_PARAGRAPH.CENTER
        scope_run = scope.add_run(f"报告范围：{meta.get('report_scope', '全部项目')}")
        scope_run.font.size = Pt(14)

        self.doc.add_paragraph()

        # 时间范围
        period = self.doc.add_paragraph()
        period.alignment = WD_ALIGN_PARAGRAPH.CENTER
        period_text = f"{meta.get('period_start', '')} 至 {meta.get('period_end', '')}"
        period_run = period.add_run(period_text)
        period_run.font.size = Pt(14)

        for _ in range(3):
            self.doc.add_paragraph()

        # 生成时间
        generated = self.doc.add_paragraph()
        generated.alignment = WD_ALIGN_PARAGRAPH.CENTER
        gen_time = meta.get('generated_at', '')
        if hasattr(gen_time, 'strftime'):
            gen_time = gen_time.strftime('%Y年%m月%d日')
        generated_run = generated.add_run(f"生成时间：{gen_time}")
        generated_run.font.size = Pt(12)

    def _add_comprehensive_cover(self, meta: Dict):
        """添加综合报告封面"""
        for _ in range(6):
            self.doc.add_paragraph()

        # 主标题
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run(meta.get('reporting_unit', '项目采购与成本管理部门'))
        title_run.font.size = Pt(26)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0, 32, 96)
        title_run.font.name = '黑体'
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

        self.doc.add_paragraph()

        # 副标题
        subtitle = self.doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.add_run(meta.get('report_title', '综合工作报告'))
        subtitle_run.font.size = Pt(22)
        subtitle_run.font.bold = True
        subtitle_run.font.name = '黑体'
        subtitle_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

        for _ in range(2):
            self.doc.add_paragraph()

        # 密级
        confidentiality = self.doc.add_paragraph()
        confidentiality.alignment = WD_ALIGN_PARAGRAPH.CENTER
        conf_run = confidentiality.add_run(f"密级：{meta.get('confidentiality_level', '内部使用')}")
        conf_run.font.size = Pt(12)
        conf_run.font.color.rgb = RGBColor(255, 0, 0)

    # ========== 目录相关方法 ==========

    def _add_table_of_contents(self):
        """添加目录（专业版）"""
        heading = self.doc.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading_run = heading.add_run('目  录')
        heading_run.font.size = Pt(18)
        heading_run.font.bold = True
        heading_run.font.name = '黑体'
        heading_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

        self.doc.add_paragraph()

        # 目录项
        toc_items = [
            '一、执行摘要',
            '二、项目概览',
            '三、采购业务分析',
            '四、合同管理分析',
            '五、付款业务分析',
            '六、结算业务分析',
            '七、归档监控',
            '八、数据完整性分析',
            '九、业务排名',
            '十、管理建议',
        ]

        for item in toc_items:
            p = self.doc.add_paragraph(item)
            p.paragraph_format.left_indent = Cm(1)

    def _add_comprehensive_toc(self):
        """添加目录（综合版）"""
        heading = self.doc.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading_run = heading.add_run('目  录')
        heading_run.font.size = Pt(20)
        heading_run.font.bold = True
        heading_run.font.name = '黑体'
        heading_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

        self.doc.add_paragraph()

        # 目录项
        toc_items = [
            '第一部分  执行摘要',
            '第二部分  组织概览',
            '第三部分  项目深度分析',
            '第四部分  采购业务综合分析',
            '第五部分  合同管理综合分析',
            '第六部分  付款业务综合分析',
            '第七部分  结算业务综合分析',
            '第八部分  财务分析',
            '第九部分  建议与展望',
        ]

        for item in toc_items:
            p = self.doc.add_paragraph(item)
            p.paragraph_format.left_indent = Cm(1.5)
            p.paragraph_format.space_after = Pt(6)

    # ========== 内容添加方法 ==========

    def _add_section_title(self, title: str, level: int = 1):
        """添加章节标题"""
        p = self.doc.add_paragraph()
        run = p.add_run(title)
        run.font.bold = True
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

        if level == 1:
            run.font.size = Pt(16)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
        elif level == 2:
            run.font.size = Pt(14)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(3)
        else:
            run.font.size = Pt(12)

    def _add_summary_content(self, summary: Dict):
        """添加摘要内容"""
        p = self.doc.add_paragraph()
        content = f"本报告统计期间为{summary.get('period_start', '')}至{summary.get('period_end', '')}。"
        content += f"期间共完成采购{summary.get('total_procurement', 0)}项，"
        content += f"签订合同{summary.get('total_contract', 0)}份，"
        content += f"处理付款{summary.get('total_payment', 0)}笔，"
        content += f"完成结算{summary.get('total_settlement', 0)}项。"
        p.add_run(content)

    def _add_procurement_section(self, data: Dict):
        """添加采购业务章节"""
        p = self.doc.add_paragraph()
        p.add_run(f"采购总数：{data.get('total_count', 0)}项")
        p = self.doc.add_paragraph()
        p.add_run(f"中标总额：{data.get('total_winning_amount', 0):,.2f}元")

    def _add_contract_section(self, data: Dict):
        """添加合同管理章节"""
        p = self.doc.add_paragraph()
        p.add_run(f"合同总数：{data.get('total_count', 0)}份")
        p = self.doc.add_paragraph()
        p.add_run(f"合同总额：{data.get('total_amount', 0):,.2f}元")

    def _add_payment_section(self, data: Dict):
        """添加付款业务章节"""
        p = self.doc.add_paragraph()
        p.add_run(f"付款总数：{data.get('total_count', 0)}笔")
        p = self.doc.add_paragraph()
        p.add_run(f"付款总额：{data.get('total_amount', 0):,.2f}元")

    def _add_settlement_section(self, data: Dict):
        """添加结算业务章节"""
        p = self.doc.add_paragraph()
        p.add_run(f"结算总数：{data.get('total_count', 0)}项")
        p = self.doc.add_paragraph()
        p.add_run(f"结算总额：{data.get('total_amount', 0):,.2f}元")

    def _add_projects_overview(self, data: Dict):
        """添加项目概览"""
        p = self.doc.add_paragraph()
        p.add_run(f"项目总数：{data.get('total_count', 0)}个")

    def _add_archive_monitoring(self, data: Dict):
        """添加归档监控"""
        p = self.doc.add_paragraph()
        p.add_run(f"归档进度：{data.get('overall_progress', 0):.1f}%")

    def _add_completeness_analysis(self, data: Dict):
        """添加完整性分析"""
        p = self.doc.add_paragraph()
        p.add_run(f"数据完整性：{data.get('overall_completeness', 0):.1f}%")

    def _add_ranking_analysis(self, data: Dict):
        """添加排名分析"""
        p = self.doc.add_paragraph()
        p.add_run("业务排名统计")

    def _add_recommendations(self, recommendations: List[str]):
        """添加管理建议"""
        for i, rec in enumerate(recommendations, 1):
            p = self.doc.add_paragraph()
            p.add_run(f"{i}. {rec}")

    def _add_organizational_overview(self, data: Dict):
        """添加组织概览（综合版）"""
        p = self.doc.add_paragraph()
        p.add_run("组织架构和职能概述")

    def _add_projects_deep_analysis(self, data: Dict):
        """添加项目深度分析（综合版）"""
        p = self.doc.add_paragraph()
        p.add_run("项目执行情况深度分析")

    def _add_business_comprehensive_sections(self, data: Dict):
        """添加业务综合分析章节（综合版）"""
        if data.get('procurement_comprehensive'):
            self._add_section_title('第四部分  采购业务综合分析', level=1)
            self._add_procurement_section(data['procurement_comprehensive'])
            self.doc.add_page_break()

    def _add_financial_analysis(self, data: Dict):
        """添加财务分析"""
        p = self.doc.add_paragraph()
        p.add_run(f"预算总额：{data.get('total_budget', 0):,.2f}元")
        p = self.doc.add_paragraph()
        p.add_run(f"节约金额：{data.get('savings_amount', 0):,.2f}元")
        p = self.doc.add_paragraph()
        p.add_run(f"节约率：{data.get('savings_rate', 0):.2f}%")
