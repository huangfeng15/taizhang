"""
Word文档构建器
提供专业的Word文档构建功能
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from typing import List, Dict, Any, Optional


class WordDocumentBuilder:
    """Word文档构建器类"""
    
    def __init__(self):
        """初始化文档构建器"""
        self.doc = Document()
        self._setup_styles()
    
    def _setup_styles(self):
        """设置文档样式"""
        # 设置默认字体为中文字体
        self.doc.styles['Normal'].font.name = 'Microsoft YaHei'
        self.doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        self.doc.styles['Normal'].font.size = Pt(10.5)
        
        # 设置标题样式
        for i in range(1, 4):
            heading_style = self.doc.styles[f'Heading {i}']
            heading_style.font.name = 'Microsoft YaHei'
            heading_style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            heading_style.font.color.rgb = RGBColor(0, 0, 0)
            heading_style.font.bold = True
    
    def create_cover_page(self, title: str, period: str, generate_date: str,
                         organization: str = ''):
        """
        创建封面页
        
        Args:
            title: 报告标题
            period: 报告周期
            generate_date: 生成日期
            organization: 单位名称
        """
        # 添加标题
        title_para = self.doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(22)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0, 51, 102)
        
        # 添加空行
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        
        # 添加周期信息
        period_para = self.doc.add_paragraph()
        period_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        period_run = period_para.add_run(f'报告周期：{period}')
        period_run.font.size = Pt(14)
        
        # 添加空行
        self.doc.add_paragraph()
        
        # 添加生成日期
        date_para = self.doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(f'生成日期：{generate_date}')
        date_run.font.size = Pt(12)
        
        # 如果有单位名称
        if organization:
            self.doc.add_paragraph()
            self.doc.add_paragraph()
            org_para = self.doc.add_paragraph()
            org_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            org_run = org_para.add_run(organization)
            org_run.font.size = Pt(14)
            org_run.font.bold = True
        
        # 添加分页符
        self.doc.add_page_break()
    
    def add_chapter(self, title: str, level: int = 1):
        """
        添加章节标题
        
        Args:
            title: 章节标题
            level: 标题级别 (1-3)
        """
        self.doc.add_heading(title, level=level)
    
    def add_paragraph(self, text: str, bold: bool = False, italic: bool = False,
                     alignment: str = 'left', font_size: int = 10.5):
        """
        添加段落
        
        Args:
            text: 段落文本
            bold: 是否加粗
            italic: 是否斜体
            alignment: 对齐方式 ('left', 'center', 'right', 'justify')
            font_size: 字体大小
        """
        para = self.doc.add_paragraph()
        run = para.add_run(text)
        run.font.bold = bold
        run.font.italic = italic
        run.font.size = Pt(font_size)
        
        # 设置对齐方式
        alignment_map = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        para.alignment = alignment_map.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)
    
    def add_bullet_list(self, items: List[str], indent_level: int = 0):
        """
        添加项目符号列表
        
        Args:
            items: 列表项
            indent_level: 缩进级别
        """
        for item in items:
            para = self.doc.add_paragraph(item, style='List Bullet')
            if indent_level > 0:
                para.paragraph_format.left_indent = Inches(0.5 * indent_level)
    
    def add_number_list(self, items: List[str], indent_level: int = 0):
        """
        添加编号列表
        
        Args:
            items: 列表项
            indent_level: 缩进级别
        """
        for item in items:
            para = self.doc.add_paragraph(item, style='List Number')
            if indent_level > 0:
                para.paragraph_format.left_indent = Inches(0.5 * indent_level)
    
    def add_table(self, data: List[List[Any]], headers: Optional[List[str]] = None,
                 col_widths: Optional[List[float]] = None, style: str = 'Light Grid Accent 1'):
        """
        添加表格
        
        Args:
            data: 表格数据（二维列表）
            headers: 表头（如果提供）
            col_widths: 列宽（单位：英寸），如果为None则自动平分
            style: 表格样式
        """
        # 确定行列数
        if headers:
            rows = len(data) + 1
            cols = len(headers)
        else:
            rows = len(data)
            cols = len(data[0]) if data else 0
        
        if cols == 0:
            return
        
        # 创建表格
        table = self.doc.add_table(rows=rows, cols=cols)
        table.style = style
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 设置列宽
        if col_widths:
            for i, width in enumerate(col_widths):
                if i < len(table.columns):
                    table.columns[i].width = Inches(width)
        else:
            # 自动平分宽度
            total_width = 6.0  # 总宽度6英寸
            col_width = total_width / cols
            for col in table.columns:
                col.width = Inches(col_width)
        
        # 填充表头
        if headers:
            header_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                cell = header_cells[i]
                cell.text = str(header)
                # 设置表头样式
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(10)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            start_row = 1
        else:
            start_row = 0
        
        # 填充数据
        for row_idx, row_data in enumerate(data):
            cells = table.rows[start_row + row_idx].cells
            for col_idx, cell_data in enumerate(row_data):
                if col_idx < len(cells):
                    cell = cells[col_idx]
                    cell.text = str(cell_data) if cell_data is not None else ''
                    # 设置单元格样式
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(9)
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    def add_key_value_table(self, data: Dict[str, Any], title: str = ''):
        """
        添加键值对表格（适合显示统计摘要）
        
        Args:
            data: 键值对数据
            title: 表格标题
        """
        if title:
            self.add_paragraph(title, bold=True, font_size=12)
        
        # 创建两列表格
        table = self.doc.add_table(rows=len(data), cols=2)
        table.style = 'Light Grid Accent 1'
        
        # 设置列宽
        table.columns[0].width = Inches(2.5)
        table.columns[1].width = Inches(3.5)
        
        # 填充数据
        for idx, (key, value) in enumerate(data.items()):
            # 键
            key_cell = table.rows[idx].cells[0]
            key_cell.text = str(key)
            for paragraph in key_cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # 值
            value_cell = table.rows[idx].cells[1]
            value_cell.text = str(value)
            for paragraph in value_cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    def add_image(self, image_path: str, width: float = 6.0, 
                 alignment: str = 'center'):
        """
        添加图片
        
        Args:
            image_path: 图片路径
            width: 图片宽度（英寸）
            alignment: 对齐方式
        """
        para = self.doc.add_paragraph()
        
        # 设置对齐方式
        alignment_map = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT
        }
        para.alignment = alignment_map.get(alignment, WD_ALIGN_PARAGRAPH.CENTER)
        
        # 添加图片
        run = para.add_run()
        run.add_picture(image_path, width=Inches(width))
    
    def add_page_break(self):
        """添加分页符"""
        self.doc.add_page_break()
    
    def add_section_break(self):
        """添加分节符"""
        self.doc.add_section()
    
    def add_horizontal_line(self):
        """添加水平分隔线"""
        para = self.doc.add_paragraph()
        para.paragraph_format.border_bottom.width = Pt(1)
        para.paragraph_format.border_bottom.color.rgb = RGBColor(192, 192, 192)
    
    def add_empty_line(self, count: int = 1):
        """
        添加空行
        
        Args:
            count: 空行数量
        """
        for _ in range(count):
            self.doc.add_paragraph()
    
    def add_summary_box(self, title: str, content: Dict[str, Any]):
        """
        添加突出显示的摘要框
        
        Args:
            title: 摘要标题
            content: 摘要内容（键值对）
        """
        # 添加标题
        title_para = self.doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.font.bold = True
        title_run.font.size = Pt(12)
        title_run.font.color.rgb = RGBColor(0, 51, 102)
        
        # 创建表格
        table = self.doc.add_table(rows=len(content), cols=2)
        table.style = 'Light Shading Accent 1'
        
        # 设置列宽
        table.columns[0].width = Inches(2.0)
        table.columns[1].width = Inches(4.0)
        
        # 填充内容
        for idx, (key, value) in enumerate(content.items()):
            # 键
            key_cell = table.rows[idx].cells[0]
            key_cell.text = str(key)
            for paragraph in key_cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(11)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            key_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            # 值
            value_cell = table.rows[idx].cells[1]
            value_cell.text = str(value)
            for paragraph in value_cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0, 102, 204)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            value_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        self.add_empty_line()
    
    def add_highlight_text(self, text: str, color: str = 'blue'):
        """
        添加突出显示的文本
        
        Args:
            text: 文本内容
            color: 颜色 ('red', 'blue', 'green', 'orange')
        """
        color_map = {
            'red': RGBColor(220, 53, 69),
            'blue': RGBColor(0, 123, 255),
            'green': RGBColor(40, 167, 69),
            'orange': RGBColor(255, 193, 7),
            'gray': RGBColor(108, 117, 125)
        }
        
        para = self.doc.add_paragraph()
        run = para.add_run(text)
        run.font.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = color_map.get(color, RGBColor(0, 123, 255))
    
    def save(self, file_path: str):
        """
        保存文档
        
        Args:
            file_path: 保存路径
        """
        self.doc.save(file_path)
    
    def get_document(self) -> Document:
        """获取文档对象"""
        return self.doc