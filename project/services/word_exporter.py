"""
Word文档导出器
用于将报告数据导出为专业格式的Word文档
"""
from typing import Dict, Any, List
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def export_to_word_professional(report_data: Dict[str, Any], file_path: str) -> str:
    """
    导出专业Word文档格式的报告
    
    Args:
        report_data: 报告数据字典
        file_path: 导出文件路径
    
    Returns:
        str: 导出的文件路径
    """
    doc = Document()
    
    # 设置文档默认样式
    _setup_document_styles(doc)
    
    # 设置页面边距
    _setup_page_margins(doc)
    
    meta = report_data.get('meta', {})
    
    # === 封面 ===
    _add_cover_page(doc, meta)
    doc.add_page_break()
    
    # === 目录 ===
    _add_table_of_contents(doc)
    doc.add_page_break()
    
    # === 一、执行摘要 ===
    _add_executive_summary(doc, report_data.get('summary', {}), meta)
    
    # === 二、项目概览 ===
    if report_data.get('projects_overview'):
        _add_projects_overview(doc, report_data['projects_overview'], meta)
    
    # === 三、采购业务分析 ===
    if report_data.get('procurement'):
        _add_procurement_analysis(doc, report_data['procurement'])
    
    # === 四、合同管理分析 ===
    if report_data.get('contract'):
        _add_contract_analysis(doc, report_data['contract'])
    
    # === 五、付款业务分析 ===
    if report_data.get('payment'):
        _add_payment_analysis(doc, report_data['payment'])
    
    # === 六、结算业务分析 ===
    if report_data.get('settlement'):
        _add_settlement_analysis(doc, report_data['settlement'])
    
    # === 七、归档监控 ===
    if report_data.get('archive_monitoring'):
        _add_archive_monitoring(doc, report_data['archive_monitoring'])
    
    # === 八、数据完整性分析 ===
    if report_data.get('completeness'):
        _add_completeness_analysis(doc, report_data['completeness'])
    
    # === 九、业务排名 ===
    if report_data.get('ranking'):
        _add_ranking_analysis(doc, report_data['ranking'])
    
    # === 十、管理建议 ===
    if report_data.get('recommendations'):
        _add_recommendations(doc, report_data['recommendations'])
    
    # === 单项目特殊内容 ===
    if report_data.get('project_details'):
        _add_project_details(doc, report_data['project_details'])
    
    # 保存文档
    doc.save(file_path)
    return file_path


def _setup_document_styles(doc):
    """设置文档样式"""
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(12)


def _setup_page_margins(doc):
    """设置页面边距"""
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)


def _add_cover_page(doc, meta: Dict):
    """添加封面"""
    for _ in range(5):
        doc.add_paragraph()
    
    # 主标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('项目采购与成本管理')
    title_run.font.size = Pt(26)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 102)
    title_run.font.name = '黑体'
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    
    # 副标题
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(meta.get('report_title', '工作报告'))
    subtitle_run.font.size = Pt(22)
    subtitle_run.font.bold = True
    subtitle_run.font.color.rgb = RGBColor(0, 51, 102)
    subtitle_run.font.name = '黑体'
    subtitle_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    
    # 项目信息
    project_info = meta.get('project_info', {})
    if project_info:
        doc.add_paragraph()
        doc.add_paragraph()
        
        if 'project_name' in project_info:
            info = doc.add_paragraph()
            info.alignment = WD_ALIGN_PARAGRAPH.CENTER
            info_run = info.add_run(f"项目名称：{project_info['project_name']}")
            info_run.font.size = Pt(16)
            info_run.font.name = '楷体'
            info_run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
        elif 'project_count' in project_info:
            info = doc.add_paragraph()
            info.alignment = WD_ALIGN_PARAGRAPH.CENTER
            info_run = info.add_run(f"涉及项目数：{project_info['project_count']}个")
            info_run.font.size = Pt(16)
            info_run.font.name = '楷体'
            info_run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
    
    # 报告期间
    doc.add_paragraph()
    doc.add_paragraph()
    period = doc.add_paragraph()
    period.alignment = WD_ALIGN_PARAGRAPH.CENTER
    period_run = period.add_run(
        f"报告期间：{meta['period_start']} 至 {meta['period_end']}"
    )
    period_run.font.size = Pt(14)
    
    # 生成日期
    doc.add_paragraph()
    gen_date = doc.add_paragraph()
    gen_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    gen_date_run = gen_date.add_run(
        f"报告生成日期：{meta['generated_at'].strftime('%Y年%m月%d日')}"
    )
    gen_date_run.font.size = Pt(12)
    gen_date_run.font.color.rgb = RGBColor(128, 128, 128)


def _add_table_of_contents(doc):
    """添加目录"""
    heading = doc.add_heading('目  录', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    toc_items = [
        '一、执行摘要',
        '二、项目概览',
        '三、采购业务分析',
        '四、合同管理分析',
        '五、付款业务分析',
        '六、结算业务分析',
        '七、归档监控',
        '八、数据完整性分析',
        '九、业务排名',
        '十、管理建议',
    ]
    
    for item in toc_items:
        p = doc.add_paragraph(item, style='List Number')
        p.paragraph_format.left_indent = Inches(0.5)


def _add_executive_summary(doc, summary: Dict, meta: Dict):
    """添加执行摘要"""
    doc.add_heading('一、执行摘要', level=1)
    
    scope_text = meta.get('report_scope', '全部项目')
    period_text = f"{meta['period_start']} 至 {meta['period_end']}"
    
    intro = doc.add_paragraph()
    intro.add_run(
        f"本报告涵盖{scope_text}在{period_text}期间的项目采购与成本管理情况。"
        f"报告期内，共涉及项目{summary.get('total_projects', 0)}个，"
        f"完成采购{summary.get('total_procurements', 0)}项，"
        f"签订合同{summary.get('total_contracts', 0)}份，"
        f"处理付款业务{summary.get('total_payments', 0)}笔，"
        f"完成结算{summary.get('total_settlements', 0)}笔。"
    )
    
    doc.add_heading('1.1 核心指标', level=2)
    
    # 创建核心指标表格
    table = doc.add_table(rows=9, cols=3)
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '指标类别'
    hdr_cells[1].text = '指标值'
    hdr_cells[2].text = '单位/说明'
    
    data = [
        ('采购预算总额', f"{summary.get('total_budget', 0):,.2f}", '万元'),
        ('采购中标总额', f"{summary.get('total_winning', 0):,.2f}", '万元'),
        ('采购节约率', f"{summary.get('savings_rate', 0):.2f}", '%'),
        ('合同签订总额', f"{summary.get('total_contract_amount', 0):,.2f}", '万元'),
        ('累计付款金额', f"{summary.get('total_payment_amount', 0):,.2f}", '万元'),
        ('付款进度', f"{summary.get('payment_rate', 0):.2f}", '%'),
        ('结算完成金额', f"{summary.get('total_settlement_amount', 0):,.2f}", '万元'),
        ('资料归档率', f"{summary.get('archive_rate', 0):.2f}", '%'),
    ]
    
    for i, (label, value, unit) in enumerate(data, start=1):
        row_cells = table.rows[i].cells
        row_cells[0].text = label
        row_cells[1].text = str(value)
        row_cells[2].text = unit


def _add_projects_overview(doc, projects_data: Dict, meta: Dict):
    """添加项目概览"""
    doc.add_heading('二、项目概览', level=1)
    
    if 'project' in projects_data:
        project = projects_data['project']
        doc.add_paragraph(f"项目编码：{project['code']}")
        doc.add_paragraph(f"项目名称：{project['name']}")
        doc.add_paragraph(f"项目负责人：{project['manager']}")
        doc.add_paragraph(f"项目状态：{project['status']}")
        
        if project.get('description'):
            doc.add_heading('2.1 项目描述', level=2)
            doc.add_paragraph(project['description'])
        
        doc.add_heading('2.2 业务统计', level=2)
        stats_text = (
            f"截至报告期末，本项目共完成采购{projects_data.get('procurement_count', 0)}项，"
            f"签订合同{projects_data.get('contract_count', 0)}份，"
            f"合同总额{projects_data.get('total_contract_amount', 0):,.2f}万元。"
        )
        doc.add_paragraph(stats_text)
    else:
        total = projects_data.get('total_count', 0)
        doc.add_paragraph(f"报告期内涉及项目共{total}个，主要项目情况如下：")
        
        projects = projects_data.get('projects', [])[:10]
        if projects:
            table = doc.add_table(rows=len(projects) + 1, cols=5)
            table.style = 'Light Grid Accent 1'
            
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '序号'
            hdr_cells[1].text = '项目编码'
            hdr_cells[2].text = '项目名称'
            hdr_cells[3].text = '项目状态'
            hdr_cells[4].text = '合同数'
            
            for i, proj in enumerate(projects, start=1):
                cells = table.rows[i].cells
                cells[0].text = str(i)
                cells[1].text = proj['code']
                cells[2].text = proj['name']
                cells[3].text = proj['status']
                cells[4].text = str(proj['contract_count'])


def _add_procurement_analysis(doc, procurement_data: Dict):
    """添加采购业务分析"""
    doc.add_heading('三、采购业务分析', level=1)
    
    stats = procurement_data
    
    doc.add_heading('3.1 采购概况', level=2)
    overview_text = (
        f"报告期内共完成采购{stats['total_count']}项，"
        f"预算总额{stats['total_budget']:,.2f}万元，"
        f"中标总额{stats['total_winning']:,.2f}万元，"
        f"节约率{stats['savings_rate']:.2f}%。"
    )
    doc.add_paragraph(overview_text)
    
    if stats.get('method_distribution'):
        doc.add_heading('3.2 采购方式分布', level=2)
        table = doc.add_table(rows=len(stats['method_distribution']) + 1, cols=4)
        table.style = 'Light Grid Accent 1'
        
        hdr = table.rows[0].cells
        hdr[0].text = '采购方式'
        hdr[1].text = '项目数'
        hdr[2].text = '中标金额（万元）'
        hdr[3].text = '占比'
        
        for i, method in enumerate(stats['method_distribution'], start=1):
            cells = table.rows[i].cells
            cells[0].text = method['method'] or '未分类'
            cells[1].text = str(method['count'])
            cells[2].text = f"{method['amount']:,.2f}"
            cells[3].text = f"{method['percentage']:.1f}%"


def _add_contract_analysis(doc, contract_data: Dict):
    """添加合同管理分析"""
    doc.add_heading('四、合同管理分析', level=1)
    
    stats = contract_data
    
    overview_text = (
        f"报告期内共签订合同{stats['total_count']}份，"
        f"合同总额{stats['total_amount']:,.2f}万元。"
        f"其中主合同{stats['main_count']}份（{stats['main_amount']:,.2f}万元），"
        f"补充协议{stats['supplement_count']}份（{stats['supplement_amount']:,.2f}万元）。"
    )
    doc.add_paragraph(overview_text)


def _add_payment_analysis(doc, payment_data: Dict):
    """添加付款业务分析"""
    doc.add_heading('五、付款业务分析', level=1)
    
    stats = payment_data
    
    overview_text = (
        f"报告期内共处理付款{stats['total_count']}笔，"
        f"付款总额{stats['total_amount']:,.2f}万元，"
        f"平均付款金额{stats['avg_amount']:,.2f}万元。"
        f"预计剩余支付金额{stats.get('estimated_remaining', 0):,.2f}万元。"
    )
    doc.add_paragraph(overview_text)


def _add_settlement_analysis(doc, settlement_data: Dict):
    """添加结算业务分析"""
    doc.add_heading('六、结算业务分析', level=1)
    
    stats = settlement_data
    
    overview_text = (
        f"报告期内共完成结算{stats['total_count']}笔，"
        f"结算总额{stats['total_amount']:,.2f}万元，"
        f"平均结算金额{stats['avg_amount']:,.2f}万元，"
        f"结算率{stats['settlement_rate']:.2f}%。"
    )
    doc.add_paragraph(overview_text)


def _add_archive_monitoring(doc, archive_data: Dict):
    """添加归档监控"""
    doc.add_heading('七、归档监控', level=1)
    
    overview = archive_data['overview']
    
    doc.add_paragraph(
        f"总体归档率：{overview['overall_rate']:.1f}%，"
        f"归档及时率：{overview.get('overall_timely_rate', 0):.1f}%。"
    )
    
    if archive_data.get('overdue_count', 0) > 0:
        doc.add_heading('7.1 逾期预警', level=2)
        doc.add_paragraph(
            f"发现逾期项目{archive_data['overdue_count']}个，"
            f"其中严重逾期{archive_data['overdue_severe']}个，"
            f"中度逾期{archive_data['overdue_moderate']}个，"
            f"轻度逾期{archive_data['overdue_mild']}个。"
        )


def _add_completeness_analysis(doc, completeness_data: Dict):
    """添加数据完整性分析"""
    doc.add_heading('八、数据完整性分析', level=1)
    
    procurement_check = completeness_data.get('procurement_field_check', {})
    contract_check = completeness_data.get('contract_field_check', {})
    
    doc.add_paragraph(
        f"采购数据完整率：{procurement_check.get('completeness_rate', 0):.1f}%，"
        f"合同数据完整率：{contract_check.get('completeness_rate', 0):.1f}%。"
    )


def _add_ranking_analysis(doc, ranking_data: Dict):
    """添加业务排名"""
    doc.add_heading('九、业务排名', level=1)
    
    doc.add_heading('9.1 采购准时完成率排名TOP5', level=2)
    on_time = ranking_data.get('procurement_on_time', [])[:5]
    if on_time:
        table = doc.add_table(rows=len(on_time) + 1, cols=4)
        table.style = 'Light Grid Accent 1'
        
        hdr = table.rows[0].cells
        hdr[0].text = '排名'
        hdr[1].text = '项目名称'
        hdr[2].text = '准时完成率'
        hdr[3].text = '采购总数'
        
        for i, item in enumerate(on_time, start=1):
            cells = table.rows[i].cells
            cells[0].text = f"{item.get('medal', '')}{i}"
            cells[1].text = item.get('name', '')
            cells[2].text = f"{item.get('on_time_rate', 0):.1f}%"
            cells[3].text = str(item.get('total_count', 0))


def _add_recommendations(doc, recommendations: List[str]):
    """添加管理建议"""
    doc.add_heading('十、管理建议', level=1)
    
    for i, recommendation in enumerate(recommendations, start=1):
        p = doc.add_paragraph(style='List Number')
        p.add_run(recommendation)


def _add_project_details(doc, project_details: Dict):
    """添加项目详细信息（单项目报告专用）"""
    doc.add_heading('附录：项目详细信息', level=1)
    
    progress = project_details.get('progress', {})
    financial = project_details.get('financial_summary', {})
    
    doc.add_heading('A.1 项目进度', level=2)
    doc.add_paragraph(
        f"采购完成数：{progress.get('procurement_count', 0)}项\n"
        f"合同签订数：{progress.get('contract_count', 0)}份\n"
        f"付款笔数：{progress.get('payment_count', 0)}笔\n"
        f"付款进度：{progress.get('payment_progress', 0):.1f}%\n"
        f"结算率：{progress.get('settlement_rate', 0):.1f}%"
    )
    
    doc.add_heading('A.2 财务汇总', level=2)
    doc.add_paragraph(
        f"合同总额：{financial.get('total_contract_amount', 0):,.2f}元\n"
        f"累计付款：{financial.get('total_paid', 0):,.2f}元\n"
        f"剩余金额：{financial.get('remaining_amount', 0):,.2f}元"
    )