"""
综合Word文档导出器
专为详细的年度总结和部门总结报告设计
支持丰富的格式、表格、图表等
"""
from typing import Dict, Any, List
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def export_comprehensive_word_report(report_data: Dict[str, Any], file_path: str) -> str:
    """
    导出综合Word报告
    
    Args:
        report_data: 报告数据（来自AdvancedReportGenerator）
        file_path: 导出文件路径
    
    Returns:
        str: 导出的文件路径
    """
    doc = Document()
    
    # 设置文档样式
    _setup_document_styles(doc)
    _setup_page_margins(doc)
    
    meta = report_data.get('meta', {})
    
    # === 封面 ===
    _add_comprehensive_cover(doc, meta)
    doc.add_page_break()
    
    # === 目录 ===
    _add_comprehensive_toc(doc)
    doc.add_page_break()
    
    # === 第一部分：执行摘要 ===
    _add_executive_summary_section(doc, report_data.get('executive_summary', {}), meta)
    doc.add_page_break()
    
    # === 第二部分：组织概览 ===
    if report_data.get('organizational_overview'):
        _add_organizational_section(doc, report_data['organizational_overview'])
        doc.add_page_break()
    
    # === 第三部分：项目分析 ===
    if report_data.get('projects_analysis'):
        _add_projects_section(doc, report_data['projects_analysis'])
        doc.add_page_break()
    
    # === 第四-七部分：业务分析 ===
    _add_business_sections(doc, report_data)
    
    # === 第八部分：建议与展望 ===
    if report_data.get('recommendations'):
        _add_recommendations_section(doc, report_data['recommendations'])
    
    # 保存文档
    doc.save(file_path)
    return file_path


def _setup_document_styles(doc):
    """设置文档样式"""
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5


def _setup_page_margins(doc):
    """设置页面边距"""
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)


def _add_comprehensive_cover(doc, meta: Dict):
    """添加综合报告封面"""
    # 顶部空白
    for _ in range(6):
        doc.add_paragraph()
    
    # 主标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(meta.get('reporting_unit', '项目采购与成本管理部门'))
    title_run.font.size = Pt(26)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 32, 96)
    title_run.font.name = '黑体'
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    
    doc.add_paragraph()
    
    # 副标题
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(meta.get('report_title', '工作总结报告'))
    subtitle_run.font.size = Pt(22)
    subtitle_run.font.bold = True
    subtitle_run.font.color.rgb = RGBColor(0, 32, 96)
    subtitle_run.font.name = '黑体'
    subtitle_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    
    # 项目信息
    doc.add_paragraph()
    doc.add_paragraph()
    
    project_info = meta.get('project_info', {})
    if project_info:
        if 'project_name' in project_info:
            info = doc.add_paragraph()
            info.alignment = WD_ALIGN_PARAGRAPH.CENTER
            info_run = info.add_run(f"项目名称：{project_info['project_name']}")
            info_run.font.size = Pt(16)
            info_run.font.name = '楷体'
            info_run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
        elif 'project_count' in project_info:
            info = doc.add_paragraph()
            info.alignment = WD_ALIGN_PARAGRAPH.CENTER
            info_run = info.add_run(f"涵盖项目：{project_info['project_count']}个")
            info_run.font.size = Pt(16)
            info_run.font.name = '楷体'
            info_run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
    
    # 报告期间
    doc.add_paragraph()
    period = doc.add_paragraph()
    period.alignment = WD_ALIGN_PARAGRAPH.CENTER
    period_run = period.add_run(
        f"报告期间：{meta['period_start']} 至 {meta['period_end']}"
    )
    period_run.font.size = Pt(14)
    
    # 生成日期
    doc.add_paragraph()
    gen_date = doc.add_paragraph()
    gen_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    gen_date_run = gen_date.add_run(
        f"报告生成日期：{meta['generated_at'].strftime('%Y年%m月%d日')}"
    )
    gen_date_run.font.size = Pt(12)
    gen_date_run.font.color.rgb = RGBColor(128, 128, 128)


def _add_comprehensive_toc(doc):
    """添加详细目录"""
    heading = doc.add_heading('目  录', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    toc_items = [
        '第一部分  执行摘要',
        '第二部分  组织概览',
        '第三部分  项目分析',
        '第四部分  采购业务分析',
        '第五部分  合同管理分析',
        '第六部分  付款管理分析',
        '第七部分  结算管理分析',
        '第八部分  建议与展望',
    ]
    
    for i, item in enumerate(toc_items, start=1):
        p = doc.add_paragraph(item, style='List Number')
        p.paragraph_format.left_indent = Inches(0.5)


def _add_executive_summary_section(doc, summary: Dict, meta: Dict):
    """添加执行摘要部分"""
    doc.add_heading('第一部分  执行摘要', level=1)
    
    # 期间描述
    period_desc = summary.get('period_description', '')
    scope_desc = summary.get('scope_description', '')
    
    intro = doc.add_paragraph()
    intro.add_run(
        f"{period_desc}期间，{scope_desc}。报告期内，各项工作稳步推进，"
        f"业务管理规范有序，取得了良好的工作成效。"
    )
    
    doc.add_heading('一、核心成就', level=2)
    
    achievements = summary.get('core_achievements', {})
    financial = summary.get('financial_summary', {})
    
    # 核心数据表格
    table = doc.add_table(rows=9, cols=3)
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '指标类别'
    hdr_cells[1].text = '数值'
    hdr_cells[2].text = '单位/说明'
    
    data = [
        ('业务数量', '', ''),
        ('  采购完成数', str(achievements.get('procurement_count', 0)), '项'),
        ('  合同签订数', str(achievements.get('contract_count', 0)), '份'),
        ('  付款笔数', str(achievements.get('payment_count', 0)), '笔'),
        ('  结算完成数', str(achievements.get('settlement_count', 0)), '笔'),
        ('财务指标', '', ''),
        ('  采购预算总额', f"{financial.get('total_budget', 0):,.2f}", '万元'),
        ('  采购节约率', f"{financial.get('savings_rate', 0):.2f}", '%'),
        ('  合同总额', f"{financial.get('total_contract', 0):,.2f}", '万元'),
        ('  累计付款', f"{financial.get('total_payment', 0):,.2f}", '万元'),
        ('  付款进度', f"{financial.get('payment_rate', 0):.2f}", '%'),
    ]
    
    for i, (label, value, unit) in enumerate(data, start=1):
        row_cells = table.rows[i].cells if i < len(table.rows) else table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = str(value)
        row_cells[2].text = unit
    
    # 工作亮点
    doc.add_heading('二、工作亮点', level=2)
    highlights = summary.get('highlights', [])
    for highlight in highlights:
        p = doc.add_paragraph(highlight, style='List Bullet')
    
    # 面临挑战
    doc.add_heading('三、面临挑战', level=2)
    challenges = summary.get('challenges', [])
    for challenge in challenges:
        p = doc.add_paragraph(challenge, style='List Bullet')


def _add_organizational_section(doc, org_data: Dict):
    """添加组织概览部分"""
    doc.add_heading('第二部分  组织概览', level=1)
    
    doc.add_heading('一、部门职责', level=2)
    doc.add_paragraph(org_data.get('mission', ''))
    
    doc.add_heading('二、核心职责', level=2)
    responsibilities = org_data.get('core_responsibilities', [])
    for resp in responsibilities:
        doc.add_paragraph(resp, style='List Number')
    
    doc.add_heading('三、工作原则', level=2)
    principles = org_data.get('work_principles', [])
    for principle in principles:
        doc.add_paragraph(principle, style='List Bullet')


def _add_projects_section(doc, projects_data: Dict):
    """添加项目分析部分"""
    doc.add_heading('第三部分  项目分析', level=1)
    
    if 'basic_info' in projects_data:
        # 单项目详细分析
        basic = projects_data['basic_info']
        doc.add_heading('一、项目基本信息', level=2)
        doc.add_paragraph(f"项目编码：{basic['code']}")
        doc.add_paragraph(f"项目名称：{basic['name']}")
        doc.add_paragraph(f"项目负责人：{basic['manager']}")
        doc.add_paragraph(f"项目状态：{basic['status']}")
        
        if basic.get('description'):
            doc.add_heading('二、项目描述', level=2)
            doc.add_paragraph(basic['description'])
        
        lifecycle = projects_data.get('lifecycle_stats', {})
        doc.add_heading('三、业务统计', level=2)
        doc.add_paragraph(
            f"累计完成采购{lifecycle.get('total_procurements', 0)}项，"
            f"签订合同{lifecycle.get('total_contracts', 0)}份，"
            f"处理付款{lifecycle.get('total_payments', 0)}笔，"
            f"完成结算{lifecycle.get('total_settlements', 0)}笔。"
        )
    else:
        # 多项目对比分析
        doc.add_heading('一、项目概况', level=2)
        total = projects_data.get('total_count', 0)
        doc.add_paragraph(f"报告期内共涉及{total}个项目。")
        
        projects = projects_data.get('projects', [])[:10]
        if projects:
            table = doc.add_table(rows=len(projects) + 1, cols=6)
            table.style = 'Light Grid Accent 1'
            
            hdr = table.rows[0].cells
            hdr[0].text = '序号'
            hdr[1].text = '项目编码'
            hdr[2].text = '项目名称'
            hdr[3].text = '采购数'
            hdr[4].text = '合同数'
            hdr[5].text = '合同额(万元)'
            
            for i, proj in enumerate(projects, start=1):
                cells = table.rows[i].cells
                cells[0].text = str(i)
                cells[1].text = proj['code']
                cells[2].text = proj['name']
                cells[3].text = str(proj['procurement_count'])
                cells[4].text = str(proj['contract_count'])
                cells[5].text = f"{proj['contract_amount']:,.2f}"


def _add_business_sections(doc, report_data: Dict):
    """添加业务分析部分（采购、合同、付款、结算）"""
    # 采购分析
    if report_data.get('procurement_comprehensive'):
        doc.add_heading('第四部分  采购业务分析', level=1)
        procurement = report_data['procurement_comprehensive']
        overview = procurement.get('overview', {})
        doc.add_paragraph(
            f"报告期内共完成采购{overview.get('total_count', 0)}项，"
            f"预算总额{overview.get('total_budget', 0):,.2f}万元，"
            f"中标总额{overview.get('total_winning', 0):,.2f}万元。"
        )
        doc.add_page_break()
    
    # 合同分析
    if report_data.get('contract_comprehensive'):
        doc.add_heading('第五部分  合同管理分析', level=1)
        contract = report_data['contract_comprehensive']
        overview = contract.get('overview', {})
        doc.add_paragraph(
            f"报告期内共签订合同{overview.get('total_count', 0)}份，"
            f"合同总额{overview.get('total_amount', 0):,.2f}万元。"
        )
        doc.add_page_break()
    
    # 付款分析
    if report_data.get('payment_comprehensive'):
        doc.add_heading('第六部分  付款管理分析', level=1)
        payment = report_data['payment_comprehensive']
        overview = payment.get('overview', {})
        doc.add_paragraph(
            f"报告期内共处理付款{overview.get('total_count', 0)}笔，"
            f"付款总额{overview.get('total_amount', 0):,.2f}万元。"
        )
        doc.add_page_break()
    
    # 结算分析
    if report_data.get('settlement_comprehensive'):
        doc.add_heading('第七部分  结算管理分析', level=1)
        settlement = report_data['settlement_comprehensive']
        overview = settlement.get('overview', {})
        doc.add_paragraph(
            f"报告期内共完成结算{overview.get('total_count', 0)}笔，"
            f"结算总额{overview.get('total_amount', 0):,.2f}万元。"
        )
        doc.add_page_break()


def _add_recommendations_section(doc, recommendations: List[Dict]):
    """添加建议与展望部分"""
    doc.add_heading('第八部分  建议与展望', level=1)
    
    doc.add_heading('一、管理建议', level=2)
    
    for i, rec in enumerate(recommendations, start=1):
        if isinstance(rec, dict):
            doc.add_heading(f'({i}) {rec.get("category", "建议")}', level=3)
            doc.add_paragraph(f"优先级：{rec.get('priority', '中')}")
            doc.add_paragraph(f"建议内容：{rec.get('description', '')}")
            if rec.get('expected_benefit'):
                doc.add_paragraph(f"预期效益：{rec['expected_benefit']}")
        else:
            doc.add_paragraph(f"{i}. {rec}", style='List Number')
    
    doc.add_heading('二、工作展望', level=2)
    doc.add_paragraph(
        "展望未来，我们将继续深化采购与成本管理工作，不断提升管理水平和工作效率，"
        "为项目建设提供更加有力的保障。主要工作重点包括："
    )
    doc.add_paragraph("持续优化采购流程，提高采购效率", style='List Bullet')
    doc.add_paragraph("加强成本控制，提升资金使用效益", style='List Bullet')
    doc.add_paragraph("强化合同管理，保障合同规范履行", style='List Bullet')
    doc.add_paragraph("完善归档机制，确保资料完整及时", style='List Bullet')
    doc.add_paragraph("提升信息化水平，推进数字化管理", style='List Bullet')
    